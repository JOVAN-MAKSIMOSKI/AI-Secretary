"""Word (docx) template renderer, shared by every docx-rendered document.

Both waste forms (identification and transport) render through this one function — it is
deliberately form-neutral: it knows only about {{token}} syntax and the values mapping it
is handed, never about columns, parties, or which form is being produced. Each route owns
that mapping and passes it in.

The only other renderer in the service is openpyxl (invoices); the offer docx route
streams its template unrendered. It replaces {{token}} placeholders at the RUN level so
every font, tab stop, and spacing setting in the template survives — the python-docx
guardrail is that `paragraph.text = ...` destroys all run formatting, so it is never
used here.

A placeholder is usually a single run (the tokenized template is authored that way), but
a token can be split across runs if the source doc fragments it. To be robust to both,
each paragraph's full text is rebuilt, tokens are substituted, and the result is written
back into the FIRST run with the remaining runs cleared — the first run's formatting is
kept for the whole paragraph. Paragraphs with no placeholder are left completely
untouched, so their multi-run formatting is preserved exactly.
"""

from __future__ import annotations

import re
from datetime import date, datetime
from io import BytesIO
from typing import Any, Mapping

from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph

# Matches {{ token }} with optional inner whitespace; the token is any run of chars that
# are not braces. Group 1 is the trimmed token key.
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")


def _format_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "true" if value else "false"
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    return str(value)


def _flatten(values: Mapping[str, Any], prefix: str = "") -> dict[str, str]:
    """Flatten nested mappings to dot-path keys, e.g. {'firm': {'name': ..}} -> firm.name.

    Matches the token style in the template ({{firm.name}}, {{contact.email}}).
    """
    flat: dict[str, str] = {}
    for key, value in values.items():
        current = f"{prefix}.{key}" if prefix else str(key)
        if isinstance(value, Mapping):
            flat.update(_flatten(value, prefix=current))
        else:
            flat[current] = _format_value(value)
    return flat


def _substitute(text: str, replacements: Mapping[str, str]) -> str:
    """Replace every {{token}} in text. An unknown token is emptied rather than left
    literally on the document, so a placeholder never leaks into a legal form."""

    def _replace(match: re.Match[str]) -> str:
        key = match.group(1)
        return replacements.get(key, "")

    return _PLACEHOLDER_RE.sub(_replace, text)


def _render_paragraph(paragraph: Paragraph, replacements: Mapping[str, str]) -> None:
    if "{{" not in paragraph.text:
        # No placeholder — leave the paragraph and all its runs untouched.
        return

    original = paragraph.text
    updated = _substitute(original, replacements)
    if updated == original:
        return

    runs = paragraph.runs
    if not runs:
        return

    # Write the whole substituted text into the first run (keeping its formatting) and
    # clear the rest. Safe because we only reach here when a placeholder was present.
    runs[0].text = updated
    for run in runs[1:]:
        run.text = ""


def _iter_cell_paragraphs(cell: _Cell):
    yield from cell.paragraphs
    # Tables can nest; recurse so a placeholder in a nested table still renders.
    for table in cell.tables:
        for row in table.rows:
            for nested_cell in row.cells:
                yield from _iter_cell_paragraphs(nested_cell)


def render_docx_bytes(
    template_bytes: bytes,
    values: Mapping[str, Any],
) -> bytes:
    """Render a docx template with the provided values.

    Supported placeholder style: {{token}} and dot paths {{firm.name}}. A token may
    appear more than once — the transport form reuses {{wasteCollected}} across its
    sections 4 and 5 — and every occurrence is replaced. Returns the rendered document
    as docx bytes; the template bytes are never mutated on disk.
    """
    document = Document(BytesIO(template_bytes))
    replacements = _flatten(values)

    for paragraph in document.paragraphs:
        _render_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in _iter_cell_paragraphs(cell):
                    _render_paragraph(paragraph, replacements)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
