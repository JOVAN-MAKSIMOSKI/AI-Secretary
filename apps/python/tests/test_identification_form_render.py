"""Layer 2 — identification-form docx renderer.

Builds a synthetic docx with {{tokens}} (so the test does not depend on the real
template being present in CI) and asserts run-level substitution: tokens are filled,
no braces leak, unknown tokens empty out, and a paragraph without any placeholder keeps
its runs untouched (the formatting-preservation guarantee).
"""

from io import BytesIO

from docx import Document

from services.docx_render import (
    _substitute,
    render_docx_bytes,
)


def _build_template() -> bytes:
    document = Document()
    document.add_paragraph("Place: {{place}} Date: {{date}}")
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("Firm: {{firm.name}} / {{firm.permit}}")
    table.cell(1, 0).paragraphs[0].add_run("Contact: {{contact.email}}")
    # A paragraph with no placeholder and multiple runs — must be left untouched.
    p = document.add_paragraph()
    p.add_run("Static ")
    p.add_run("multi-run ")
    p.add_run("text")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(p.text for p in cell.paragraphs)
    return "\n".join(parts)


def test_substitute_basic() -> None:
    assert _substitute("Hello {{x}} world", {"x": "ACME"}) == "Hello ACME world"
    assert _substitute("{{a}}{{b}}", {"a": "1", "b": "2"}) == "12"


def test_substitute_unknown_token_emptied() -> None:
    assert _substitute("a {{zzz}} b", {}) == "a  b"


def test_render_fills_tokens_in_body_and_tables() -> None:
    rendered = render_docx_bytes(
        _build_template(),
        {
            "place": "Skopje",
            "date": "2026-07-25",
            "firm": {"name": "ACME DOO", "permit": "PERMIT-9"},
            "contact": {"email": "jovan@acme.mk"},
        },
    )
    text = _all_text(Document(BytesIO(rendered)))

    for value in ("Skopje", "2026-07-25", "ACME DOO", "PERMIT-9", "jovan@acme.mk"):
        assert value in text
    assert "{{" not in text and "}}" not in text


def test_render_preserves_placeholderless_paragraph_runs() -> None:
    rendered = render_docx_bytes(
        _build_template(),
        {"place": "Skopje", "date": "2026-07-25", "firm": {}, "contact": {}},
    )
    doc = Document(BytesIO(rendered))
    static = [p for p in doc.paragraphs if p.text == "Static multi-run text"]
    assert static, "the placeholder-free paragraph should be unchanged"
    # Its multiple runs must survive (not collapsed into one), proving we skipped it.
    assert len(static[0].runs) == 3


def test_render_does_not_mutate_input_bytes() -> None:
    template = _build_template()
    original = bytes(template)
    render_docx_bytes(template, {"place": "X", "date": "2026-01-01", "firm": {}, "contact": {}})
    assert template == original
