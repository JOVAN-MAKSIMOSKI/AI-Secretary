"""Layer 3 — /documents/identification-form route logic.

Drives the real handler through FastAPI's TestClient with the supabase client, template
fetch, and upload all stubbed (no network). Guards the route's job: resolve + verify the
parties, reject a form whose required legal boxes would be blank, then insert/render/
upload on the happy path. These are the checks the Pydantic model cannot make.
"""

from io import BytesIO

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

import routers.documents as rd
from models.documents import (
    _HAZARDOUS_EWC_CODES,
    _PACKING_METHODS,
    _WASTE_ORIGINS,
    _WASTE_OPERATIONS_CODES,
)
from services.auth import get_current_user_id_or_service

_TENANT = "8da53b32-c30c-45bc-a148-49f468bda999"


def _template_bytes() -> bytes:
    document = Document()
    document.add_paragraph(
        "{{firm.name}} {{firm.permit}} {{firm.address}} {{contact.email}} "
        "{{wasteCode}} {{place}} {{date}}"
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class _Query:
    def __init__(self, store, table, firm_row, contact_row):
        self._store = store
        self._table = table
        self._firm_row = firm_row
        self._contact_row = contact_row
        self._filters = {}

    def select(self, *args, **kwargs):
        return self

    def eq(self, column, value):
        self._filters[column] = value
        return self

    def limit(self, *args, **kwargs):
        return self

    def order(self, *args, **kwargs):
        return self

    def insert(self, data):
        self._store.setdefault("inserted", []).append(data)
        return self

    def delete(self):
        self._store["deleted"] = True
        return self

    def execute(self):
        class _Response:
            pass

        response = _Response()
        if self._table == "businesses":
            response.data = [{"owner_auth_id": _TENANT}]
        elif self._table == "firms":
            response.data = [self._firm_row] if self._firm_row else []
        elif self._table == "contacts":
            response.data = [self._contact_row] if self._contact_row else []
        else:
            response.data = []
        return response


class _Supabase:
    def __init__(self, store, firm_row, contact_row):
        self._store = store
        self._firm_row = firm_row
        self._contact_row = contact_row

    def table(self, name):
        return _Query(self._store, name, self._firm_row, self._contact_row)


def _client(monkeypatch, firm_row, contact_row, store):
    monkeypatch.setattr(rd, "supabase", _Supabase(store, firm_row, contact_row))
    monkeypatch.setattr(
        rd,
        "fetch_identification_form_template_payload",
        lambda tenant_id: {"template_bytes": _template_bytes(), "template_name": "idf"},
    )
    monkeypatch.setattr(rd, "upload_identification_form_document", lambda *a, **k: store.setdefault("uploads", []).append(a))

    app = FastAPI()
    app.include_router(rd.router)
    app.dependency_overrides[get_current_user_id_or_service] = lambda: _TENANT
    return TestClient(app)


def _payload(**overrides):
    payload = {
        "firm_id": "firm-1",
        "contact_id": "contact-1",
        "firm_name": "ACME DOO",
        "waste_location": "Skopje warehouse",
        "is_hazardous": True,
        "waste_type": "Опасни агрохемиски отпадоци",
        "ewc_code": next(iter(_HAZARDOUS_EWC_CODES)),
        "packing_method": next(iter(_PACKING_METHODS)),
        "total_weight_kg": "1500.00",
        "waste_origin": next(iter(_WASTE_ORIGINS)),
        "waste_operation_code": next(iter(_WASTE_OPERATIONS_CODES)),
        "place": "Skopje",
        "date": "2026-07-25",
    }
    payload.update(overrides)
    return payload


_GOOD_FIRM = {"id": "firm-1", "name": "ACME DOO", "address": "ul. 1", "permit_number": "P-9"}
_GOOD_CONTACT = {
    "id": "contact-1",
    "name": "Jovan",
    "phone_number": "070",
    "email": "jovan@acme.mk",
    "firm_id": "firm-1",
}


def test_happy_path_renders_and_persists(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, _GOOD_FIRM, _GOOD_CONTACT, store)

    response = client.post("/documents/identification-form", json=_payload())

    assert response.status_code == 200, response.text
    assert response.headers["content-type"] == rd._DOCX_MEDIA_TYPE
    text = Document(BytesIO(response.content)).paragraphs[0].text
    assert "ACME DOO" in text and "P-9" in text and "jovan@acme.mk" in text
    assert "{{" not in text
    inserted = store["inserted"][0]
    assert inserted["storagePath"] == f"{_TENANT}/identification-forms/{inserted['id']}.docx"
    assert store.get("uploads")


def test_unknown_firm_404(monkeypatch) -> None:
    client = _client(monkeypatch, None, _GOOD_CONTACT, {})
    response = client.post("/documents/identification-form", json=_payload())
    assert response.status_code == 404


def test_firm_name_mismatch_422(monkeypatch) -> None:
    client = _client(monkeypatch, _GOOD_FIRM, _GOOD_CONTACT, {})
    response = client.post("/documents/identification-form", json=_payload(firm_name="WRONG NAME"))
    assert response.status_code == 422


def test_firm_missing_address_422(monkeypatch) -> None:
    firm = {**_GOOD_FIRM, "address": ""}
    client = _client(monkeypatch, firm, _GOOD_CONTACT, {})
    response = client.post("/documents/identification-form", json=_payload())
    assert response.status_code == 422
    assert "address" in response.text


def test_firm_missing_permit_422(monkeypatch) -> None:
    firm = {**_GOOD_FIRM, "permit_number": None}
    client = _client(monkeypatch, firm, _GOOD_CONTACT, {})
    response = client.post("/documents/identification-form", json=_payload())
    assert response.status_code == 422
    assert "permit" in response.text


def test_contact_from_other_firm_422(monkeypatch) -> None:
    contact = {**_GOOD_CONTACT, "firm_id": "firm-2"}
    client = _client(monkeypatch, _GOOD_FIRM, contact, {})
    response = client.post("/documents/identification-form", json=_payload())
    assert response.status_code == 422
    assert "does not belong" in response.text


def test_default_origin_is_manual(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, _GOOD_FIRM, _GOOD_CONTACT, store)
    client.post("/documents/identification-form", json=_payload())
    assert store["inserted"][0]["origin"] == "manual"


def test_call_origin_header_stamps_call(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, _GOOD_FIRM, _GOOD_CONTACT, store)
    client.post(
        "/documents/identification-form",
        json=_payload(),
        headers={"X-Document-Origin": "call"},
    )
    assert store["inserted"][0]["origin"] == "call"


def test_unknown_origin_header_falls_back_to_manual(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, _GOOD_FIRM, _GOOD_CONTACT, store)
    client.post(
        "/documents/identification-form",
        json=_payload(),
        headers={"X-Document-Origin": "bogus"},
    )
    assert store["inserted"][0]["origin"] == "manual"
