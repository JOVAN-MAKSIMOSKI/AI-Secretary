"""Layer 3 — /documents/transport-form route logic.

Drives the real handler through FastAPI's TestClient with the supabase client, template
fetch, and upload all stubbed (no network). Guards the route's job: resolve + verify both
parties, pick the collector permit by is_hazardous, reject a form whose required legal
boxes would be blank, then insert/render/upload on the happy path. These are the checks
the Pydantic model cannot make.
"""

from io import BytesIO

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

import routers.documents as rd
from models.documents import _HAZARDOUS_EWC_CODES, _NON_HAZARDOUS_EWC_CODES
from services.auth import get_current_user_id_or_service

_TENANT = "8da53b32-c30c-45bc-a148-49f468bda999"
_HAZ_CODE = next(iter(_HAZARDOUS_EWC_CODES))
_NON_HAZ_CODE = next(iter(_NON_HAZARDOUS_EWC_CODES))


def _template_bytes() -> bytes:
    """Synthetic template carrying all 15 tokens of the real one, so the test does not
    depend on TransportFormTemplate.tokenized.docx being present in CI."""
    document = Document()
    document.add_paragraph(
        "{{wasteDescription}} {{wasteCode}} {{TotalWasteWeight}} "
        "{{firm.name}} {{firm.address}} "
        "{{wasteCollected}} {{wasteCollectedDate}} {{wasteCollectedPlace}} "
        "{{tennant.permitNumber}} "
        "{{disposalPlace.name}} {{disposalPlace.address}} {{disposalPlace.location}} "
        "{{totalWasteWeightDisposed}} {{dateOfDisposal}} {{note}}"
    )
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class _Query:
    def __init__(self, store, table, rows_by_table):
        self._store = store
        self._table = table
        self._rows_by_table = rows_by_table

    def select(self, *args, **kwargs):
        return self

    def eq(self, column, value):
        return self

    def limit(self, *args, **kwargs):
        return self

    def insert(self, data):
        self._store.setdefault("inserted", []).append(data)
        return self

    def delete(self):
        self._store["deleted"] = True
        return self

    def execute(self):
        class _Response:
            pass

        response = _Response()
        row = self._rows_by_table.get(self._table)
        response.data = [row] if row else []
        return response


class _Supabase:
    def __init__(self, store, rows_by_table):
        self._store = store
        self._rows_by_table = rows_by_table

    def table(self, name):
        return _Query(self._store, name, self._rows_by_table)


_GOOD_BUSINESS = {
    "owner_auth_id": _TENANT,
    "permit_number": "P-ORDINARY",
    "dangerous_waste_permit_number": "P-DANGEROUS",
}
_GOOD_FIRM = {"id": "firm-1", "name": "ACME DOO", "address": "ul. 1", "city": "Skopje"}
_GOOD_DISPOSAL = {
    "id": "dp-1",
    "name": "Drisla",
    "address": "Drisla bb",
    "place": "Kumanovo",
}


def _client(monkeypatch, store, *, business=None, firm=None, disposal=None):
    rows = {
        "businesses": _GOOD_BUSINESS if business is None else business,
        "firms": _GOOD_FIRM if firm is None else firm,
        "disposal_places": _GOOD_DISPOSAL if disposal is None else disposal,
    }
    monkeypatch.setattr(rd, "supabase", _Supabase(store, rows))
    monkeypatch.setattr(
        rd,
        "fetch_transport_form_template_payload",
        lambda tenant_id: {"template_bytes": _template_bytes(), "template_name": "tf"},
    )
    monkeypatch.setattr(
        rd,
        "upload_transport_form_document",
        lambda *a, **k: store.setdefault("uploads", []).append(a),
    )

    app = FastAPI()
    app.include_router(rd.router)
    app.dependency_overrides[get_current_user_id_or_service] = lambda: _TENANT
    return TestClient(app)


def _payload(**overrides):
    payload = {
        "firm_id": "firm-1",
        "disposal_place_id": "dp-1",
        "firm_name": "ACME DOO",
        "disposal_place_name": "Drisla",
        "waste_type": "Отпадно масло",
        "is_hazardous": True,
        "ewc_code": _HAZ_CODE,
        "waste_owner_total_kg": "900.00",
        "collector_total_kg": "800.00",
        "collector_date": "2026-03-12",
        "end_owner_total_kg": "780.00",
        "end_owner_date": "2026-03-13",
        "note": "Handle with care",
    }
    payload.update(overrides)
    return payload


def _rendered_text(response) -> str:
    return Document(BytesIO(response.content)).paragraphs[0].text


def test_happy_path_renders_and_persists(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, store)

    response = client.post("/documents/transport-form", json=_payload())

    assert response.status_code == 200, response.text
    assert response.headers["content-type"] == rd._DOCX_MEDIA_TYPE

    text = _rendered_text(response)
    assert "{{" not in text, "a placeholder leaked into the document"
    # All three weights land in their own boxes and are not conflated.
    assert "900.00" in text and "800.00" in text and "780.00" in text
    assert "ACME DOO" in text and "Drisla" in text
    assert "Handle with care" in text

    inserted = store["inserted"][0]
    assert inserted["storagePath"] == f"{_TENANT}/transport-forms/{inserted['id']}.docx"
    assert inserted["status"] == "draft"
    assert "waste_owner_date" not in inserted
    assert store.get("uploads")


def test_handover_town_comes_from_firm_city(monkeypatch) -> None:
    # Section 4/5's "Во …" is the firm's city; section 6's is the disposal place's town.
    client = _client(monkeypatch, {})
    response = client.post("/documents/transport-form", json=_payload())
    text = _rendered_text(response)
    assert "Skopje" in text and "Kumanovo" in text


def test_hazardous_uses_dangerous_waste_permit(monkeypatch) -> None:
    client = _client(monkeypatch, {})
    response = client.post("/documents/transport-form", json=_payload())
    text = _rendered_text(response)
    assert "P-DANGEROUS" in text
    assert "P-ORDINARY" not in text


def test_non_hazardous_uses_ordinary_permit(monkeypatch) -> None:
    client = _client(monkeypatch, {})
    response = client.post(
        "/documents/transport-form",
        json=_payload(is_hazardous=False, ewc_code=_NON_HAZ_CODE),
    )
    text = _rendered_text(response)
    assert "P-ORDINARY" in text
    assert "P-DANGEROUS" not in text


def test_missing_dangerous_permit_422(monkeypatch) -> None:
    business = {**_GOOD_BUSINESS, "dangerous_waste_permit_number": ""}
    client = _client(monkeypatch, {}, business=business)
    response = client.post("/documents/transport-form", json=_payload())
    assert response.status_code == 422
    assert "dangerous-waste permit" in response.text


def test_missing_ordinary_permit_422(monkeypatch) -> None:
    business = {**_GOOD_BUSINESS, "permit_number": None}
    client = _client(monkeypatch, {}, business=business)
    response = client.post(
        "/documents/transport-form",
        json=_payload(is_hazardous=False, ewc_code=_NON_HAZ_CODE),
    )
    assert response.status_code == 422
    assert "permit number" in response.text


def test_unknown_tenant_404(monkeypatch) -> None:
    client = _client(monkeypatch, {}, business={})
    response = client.post("/documents/transport-form", json=_payload())
    assert response.status_code == 404


def test_unknown_firm_404(monkeypatch) -> None:
    client = _client(monkeypatch, {}, firm={})
    response = client.post("/documents/transport-form", json=_payload())
    assert response.status_code == 404


def test_firm_name_mismatch_422(monkeypatch) -> None:
    client = _client(monkeypatch, {})
    response = client.post("/documents/transport-form", json=_payload(firm_name="WRONG"))
    assert response.status_code == 422
    assert "firm_name" in response.text


def test_firm_missing_address_422(monkeypatch) -> None:
    client = _client(monkeypatch, {}, firm={**_GOOD_FIRM, "address": ""})
    response = client.post("/documents/transport-form", json=_payload())
    assert response.status_code == 422
    assert "address" in response.text


def test_firm_missing_city_422(monkeypatch) -> None:
    # city is nullable and renders as the handover town — a blank one would leave
    # "Во ___" empty on a legal document.
    client = _client(monkeypatch, {}, firm={**_GOOD_FIRM, "city": None})
    response = client.post("/documents/transport-form", json=_payload())
    assert response.status_code == 422
    assert "city" in response.text


def test_unknown_disposal_place_404(monkeypatch) -> None:
    client = _client(monkeypatch, {}, disposal={})
    response = client.post("/documents/transport-form", json=_payload())
    assert response.status_code == 404


def test_disposal_place_name_mismatch_422(monkeypatch) -> None:
    client = _client(monkeypatch, {})
    response = client.post(
        "/documents/transport-form", json=_payload(disposal_place_name="WRONG")
    )
    assert response.status_code == 422
    assert "disposal_place_name" in response.text


def test_row_rolled_back_when_render_fails(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, store)
    monkeypatch.setattr(
        rd, "render_docx_bytes", lambda *a, **k: (_ for _ in ()).throw(RuntimeError("boom"))
    )

    response = client.post("/documents/transport-form", json=_payload())

    assert response.status_code == 502
    assert store.get("deleted") is True, "inserted row must be rolled back"


def test_row_rolled_back_when_upload_fails(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, store)
    monkeypatch.setattr(
        rd,
        "upload_transport_form_document",
        lambda *a, **k: (_ for _ in ()).throw(RuntimeError("boom")),
    )

    response = client.post("/documents/transport-form", json=_payload())

    assert response.status_code == 502
    assert store.get("deleted") is True, "inserted row must be rolled back"


def test_blank_note_renders_empty_not_placeholder(monkeypatch) -> None:
    client = _client(monkeypatch, {})
    response = client.post("/documents/transport-form", json=_payload(note="  "))
    assert response.status_code == 200
    assert "{{note}}" not in _rendered_text(response)


def test_default_origin_is_manual(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, store)
    client.post("/documents/transport-form", json=_payload())
    assert store["inserted"][0]["origin"] == "manual"


def test_call_origin_header_stamps_call(monkeypatch) -> None:
    # The Twilio path sends this header; origin='call' + downloaded_at IS NULL is what
    # the dashboard's pending-call transport-forms card queries.
    store: dict = {}
    client = _client(monkeypatch, store)
    client.post(
        "/documents/transport-form",
        json=_payload(),
        headers={"X-Document-Origin": "call"},
    )
    assert store["inserted"][0]["origin"] == "call"


def test_unknown_origin_header_falls_back_to_manual(monkeypatch) -> None:
    store: dict = {}
    client = _client(monkeypatch, store)
    client.post(
        "/documents/transport-form",
        json=_payload(),
        headers={"X-Document-Origin": "bogus"},
    )
    assert store["inserted"][0]["origin"] == "manual"
