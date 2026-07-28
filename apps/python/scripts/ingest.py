#!/usr/bin/env python3
"""Ingest Macedonian waste-law documents (PDF/DOCX) into Qdrant.

Pipeline (waste-law RAG advisor, Phase 1):
  1. Extract text — fitz for PDFs; pages with no selectable text are OCR'd
     via pytesseract (lang=mkd). DOCX files are read with python-docx.
  2. Chunk at article boundaries ("Член N"), not fixed character windows.
     Articles longer than the e5 512-token context are sub-split with
     sentence-aware overlap, keeping the same article metadata.
  3. Per-article metadata (title, waste_types, entity_types) extracted once
     at ingest time with a lightweight LLM (GitHub Models gpt-4o-mini);
     regex fallback if the LLM is unavailable.
  4. Embed locally with intfloat/multilingual-e5-large (1024-dim).
     E5 requires the "passage: " prefix on documents — queries must use
     "query: " (handled in services/ragagent.py).

Re-ingesting is non-destructive by default: point ids are derived from the
source file, article and chunk text, so a second run of the same document
overwrites its own points in place instead of duplicating them. Nothing is
deleted unless --recreate is passed, and that path snapshots first.

Usage:
    python scripts/ingest.py <file.pdf> [<file2.docx> ...]   # law/date auto-detected per file
    python scripts/ingest.py --law-name "Law on Waste Management 216/2021" \
        --valid-from 2021-01-01 <file.pdf>                   # explicit override
    python scripts/ingest.py --recreate <file.pdf> ...       # drop the collection first (snapshots)
    python scripts/ingest.py --no-metadata-llm ...           # regex-only metadata

Required env vars (loaded from apps/python/.env):
    QDRANT_LOCAL_PATH     local Qdrant storage directory (or QDRANT_URL)
    QDRANT_COLLECTION     target collection (default: waste_management_law_mk_v2)
    GITHUB_MODELS_TOKEN   only needed for LLM metadata extraction

System requirements for OCR of scanned pages:
    Tesseract binary with the Macedonian language pack ("mkd").
    Set TESSERACT_CMD if tesseract.exe is not on PATH.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import sys
import uuid
from datetime import date
from pathlib import Path

_ENV_PATH = Path(__file__).resolve().parents[1] / ".env"
if _ENV_PATH.exists():
    from dotenv import load_dotenv
    load_dotenv(_ENV_PATH)

import fitz  # PyMuPDF
from llama_index.core.node_parser import SentenceSplitter
from qdrant_client import QdrantClient
from qdrant_client.models import (
    Distance,
    FieldCondition,
    Filter,
    FilterSelector,
    MatchValue,
    PointStruct,
    VectorParams,
)

EMBED_MODEL_NAME = "intfloat/multilingual-e5-large"
EMBED_DIMS = 1024
E5_PASSAGE_PREFIX = "passage: "  # e5 documents must be embedded with this prefix
DEFAULT_COLLECTION = "waste_management_law_mk_v2"

# e5 has a 512-token window; Cyrillic tokenizes densely, so cap sub-chunks
# well below it (token counts here use SentenceSplitter's tokenizer, which
# over-counts Cyrillic vs XLM-R — erring on the safe side).
SUBCHUNK_TOKENS = 400
SUBCHUNK_OVERLAP = 50
LONG_ARTICLE_CHAR_THRESHOLD = 1200

ARTICLE_BOUNDARY_RE = re.compile(r"(?=Член\s+\d+)")
ARTICLE_NUMBER_RE = re.compile(r"^Член\s+(\d+)")

# Below this a PDF page is treated as scanned and OCR'd. Set above header/footer
# size: the corpus has scanned gazettes whose only text layer is a ~65-char header.
MIN_SELECTABLE_CHARS_PER_PAGE = 100

# Files whose extracted text is mostly non-Cyrillic are almost certainly in the
# legacy Macedonian font-hack encoding (Latin glyphs remapped by the font) —
# embedding those produces garbage vectors, so they are skipped with a warning.
MIN_CYRILLIC_RATIO = 0.3
OCR_RENDER_DPI = 300
OCR_LANG = "mkd"

METADATA_LLM_MODEL = "gpt-4o-mini"
GITHUB_MODELS_BASE_URL = "https://models.inference.ai.azure.com"

EMBED_BATCH_SIZE = 32
UPSERT_BATCH_SIZE = 100

# Fixed namespace for uuid5 point ids. Must never change: it is what makes a
# re-ingest overwrite the previous run's points instead of duplicating them.
POINT_ID_NAMESPACE = uuid.uuid5(uuid.NAMESPACE_DNS, "waste-law-ingest.ai-secretary")
CHUNK_HASH_PREFIX_LEN = 16

_REPO_ROOT = Path(__file__).resolve().parents[3]
_DEFAULT_QDRANT_PATH = _REPO_ROOT / "apps" / "agent" / "src" / "rag-agent" / "qdrant_data"

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


# ---------------------------------------------------------------------------
# Clients
# ---------------------------------------------------------------------------

def _build_qdrant_client() -> QdrantClient:
    qdrant_url = os.getenv("QDRANT_URL", "").strip()
    qdrant_api_key = os.getenv("QDRANT_API_KEY", "").strip()
    if qdrant_url:
        return QdrantClient(url=qdrant_url, api_key=qdrant_api_key or None)
    local_path = Path(os.getenv("QDRANT_LOCAL_PATH", str(_DEFAULT_QDRANT_PATH))).resolve()
    local_path.mkdir(parents=True, exist_ok=True)
    return QdrantClient(path=str(local_path))


def _build_metadata_llm_client():
    """OpenAI-compatible client for GitHub Models, or None if no token."""
    token = os.getenv("GITHUB_MODELS_TOKEN", os.getenv("RAG_GITHUB_MODELS_TOKEN", "")).strip()
    if not token:
        return None
    from openai import OpenAI
    return OpenAI(api_key=token, base_url=GITHUB_MODELS_BASE_URL)


def _snapshot_before_destroy(client: QdrantClient, collection_name: str) -> None:
    """Snapshot a collection before dropping it, or abort the run.

    Aborting is deliberate: --recreate is otherwise unrecoverable, and an
    interrupted rebuild behind it is what destroyed the corpus twice before.
    Snapshots are a server feature — embedded local-path mode has no such API,
    which is precisely why those losses had nothing to roll back to.
    """
    try:
        snapshot = client.create_snapshot(collection_name=collection_name)
    except Exception as exc:
        sys.exit(
            f"[ingest] Refusing to --recreate '{collection_name}': snapshot failed ({exc}).\n"
            "[ingest] Snapshots require a Qdrant server (QDRANT_URL). Start the dev server with\n"
            "[ingest]   docker compose -f docker-compose.dev.yml up -d"
        )
    print(f"[ingest] Snapshot taken before recreate: {getattr(snapshot, 'name', snapshot)}")


def _ensure_collection(client: QdrantClient, collection_name: str, *, recreate: bool) -> None:
    existing = {c.name for c in client.get_collections().collections}
    if collection_name in existing:
        if not recreate:
            print(f"[ingest] Upserting into existing collection '{collection_name}'.")
            return
        _snapshot_before_destroy(client, collection_name)
        print(f"[ingest] --recreate: dropping '{collection_name}' for a fresh index.")
        client.delete_collection(collection_name)
    client.create_collection(
        collection_name=collection_name,
        vectors_config=VectorParams(size=EMBED_DIMS, distance=Distance.COSINE),
    )
    print(f"[ingest] Collection '{collection_name}' created ({EMBED_DIMS}-dim Cosine).")


# ---------------------------------------------------------------------------
# Text extraction (PDF with OCR fallback, DOCX)
# ---------------------------------------------------------------------------

def _ocr_page(page: "fitz.Page") -> str:
    """Rasterize a scanned page with fitz and OCR it with Tesseract.

    Uses fitz's own renderer instead of pdf2image, which avoids the
    poppler system dependency on Windows.
    """
    import io

    import pytesseract
    from PIL import Image

    tesseract_cmd = os.getenv("TESSERACT_CMD", "").strip()
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    pixmap = page.get_pixmap(dpi=OCR_RENDER_DPI)
    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
    return pytesseract.image_to_string(image, lang=OCR_LANG)


def _extract_pdf_text(pdf_path: Path) -> str:
    doc = fitz.open(str(pdf_path))
    pages: list[str] = []
    ocr_pages = 0
    for page_index in range(len(doc)):
        page = doc[page_index]
        text = page.get_text().strip()
        if len(text) < MIN_SELECTABLE_CHARS_PER_PAGE:
            # Scanned page — no usable text layer. Official Gazette PDFs mix
            # selectable and scanned pages, so this is checked per page.
            try:
                text = _ocr_page(page).strip()
                ocr_pages += 1
            except Exception as exc:  # noqa: BLE001 — surface OCR setup problems clearly
                print(
                    f"[ingest]   ! OCR failed on page {page_index + 1} of {pdf_path.name}: {exc}\n"
                    "[ingest]     Install Tesseract with the 'mkd' language pack, "
                    "or set TESSERACT_CMD to the tesseract.exe path.",
                    file=sys.stderr,
                )
        pages.append(text)
    if ocr_pages:
        print(f"[ingest]   -> OCR'd {ocr_pages}/{len(doc)} scanned pages")
    return "\n\n".join(pages)


def _extract_docx_text(docx_path: Path) -> str:
    """Read a DOCX: body paragraphs plus table cell text, in document order."""
    from docx import Document

    doc = Document(str(docx_path))
    parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            # python-docx yields a merged cell once per spanned column —
            # drop consecutive repeats to avoid N-fold duplicated text
            cells: list[str] = []
            previous = None
            for cell in row.cells:
                text = cell.text.strip()
                if text and text != previous:
                    cells.append(text)
                previous = text
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        return _extract_pdf_text(path)
    if path.suffix.lower() == ".docx":
        return _extract_docx_text(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


# ---------------------------------------------------------------------------
# Article-level chunking
# ---------------------------------------------------------------------------

def _split_articles(full_text: str) -> list[dict]:
    """Split on 'Член N' boundaries. Text before the first article (title,
    preamble) becomes one chunk with article=None."""
    segments = [s.strip() for s in ARTICLE_BOUNDARY_RE.split(full_text) if s.strip()]
    articles: list[dict] = []
    for segment in segments:
        match = ARTICLE_NUMBER_RE.match(segment)
        articles.append({
            "text": segment,
            "article": match.group(1) if match else None,
        })
    return articles


def _subsplit_long_article(text: str) -> list[str]:
    """e5-large truncates past 512 tokens, so long articles are sub-split
    with sentence-aware overlap. Each part keeps the parent article metadata."""
    if len(text) <= LONG_ARTICLE_CHAR_THRESHOLD:
        return [text]
    splitter = SentenceSplitter(chunk_size=SUBCHUNK_TOKENS, chunk_overlap=SUBCHUNK_OVERLAP)
    return splitter.split_text(text)


# ---------------------------------------------------------------------------
# LLM metadata extraction (per article, at ingest time)
# ---------------------------------------------------------------------------

# After this many consecutive LLM failures (usually rate limits), stop calling
# the LLM for the rest of the run instead of paying a failed HTTP roundtrip per article.
METADATA_LLM_MAX_CONSECUTIVE_FAILURES = 3

_DOC_INFO_PROMPT = """From the first page of a North Macedonian legal document (Macedonian Cyrillic), extract:

Return ONLY a JSON object:
- "law": official short title + gazette number if visible, e.g. "Закон за управување со отпадот 216/2021" (empty string if not identifiable)
- "valid_from": ISO date (YYYY-MM-DD) the act was published/took effect, or null if not visible

First page text:
"""


def _detect_document_info(llm_client, first_page_text: str, fallback_name: str) -> dict:
    """Identify which law/amendment a file contains from its opening text.

    The corpus is a mixed bundle (base law + amendments + bylaws) with opaque
    filenames, so one CLI-supplied law name for all files would be wrong.
    """
    fallback = {"law": fallback_name, "valid_from": None}
    if llm_client is None or not first_page_text.strip():
        return fallback
    try:
        response = llm_client.chat.completions.create(
            model=METADATA_LLM_MODEL,
            messages=[{"role": "user", "content": _DOC_INFO_PROMPT + first_page_text[:4000]}],
            response_format={"type": "json_object"},
            temperature=0,
        )
        parsed = json.loads(response.choices[0].message.content)
        law = str(parsed.get("law", "")).strip()
        valid_from = parsed.get("valid_from")
        if isinstance(valid_from, str):
            try:
                date.fromisoformat(valid_from)
            except ValueError:
                valid_from = None
        else:
            valid_from = None
        return {"law": law or fallback_name, "valid_from": valid_from}
    except Exception as exc:  # noqa: BLE001 — detection is best-effort
        print(f"[ingest]   ! doc-info LLM failed ({exc}) — using filename as law name", file=sys.stderr)
        return fallback


_METADATA_PROMPT = """You extract metadata from one article of a North Macedonian waste-management law (Macedonian Cyrillic text).

Return ONLY a JSON object with these keys:
- "title": short title/subject of the article in English (max 10 words)
- "waste_types": array from ["hazardous", "construction", "packaging", "electronic", "municipal", "paper_textile", "general"] — only types this article concerns; [] if none specifically
- "entity_types": array from ["individual", "business", "municipality", "government"] — who the article applies to; [] if unclear

Article text:
"""


def _regex_fallback_metadata(article_text: str) -> dict:
    # First non-"Член N" line usually carries the article's subject.
    lines = [l.strip() for l in article_text.splitlines() if l.strip()]
    title = ""
    for line in lines:
        if not ARTICLE_NUMBER_RE.match(line):
            title = line[:120]
            break
    return {"title": title, "waste_types": [], "entity_types": []}


def _extract_article_metadata(llm_client, article_text: str, llm_state: dict) -> dict:
    # Circuit breaker: GitHub Models free-tier rate limits would otherwise turn
    # a large ingest into hundreds of slow failed calls.
    if llm_client is None or llm_state.get("disabled"):
        return _regex_fallback_metadata(article_text)
    try:
        response = llm_client.chat.completions.create(
            model=METADATA_LLM_MODEL,
            messages=[{"role": "user", "content": _METADATA_PROMPT + article_text[:6000]}],
            response_format={"type": "json_object"},
            temperature=0,
        )
        parsed = json.loads(response.choices[0].message.content)
        llm_state["consecutive_failures"] = 0
        return {
            "title": str(parsed.get("title", ""))[:200],
            "waste_types": [str(w) for w in parsed.get("waste_types", []) if isinstance(w, str)],
            "entity_types": [str(e) for e in parsed.get("entity_types", []) if isinstance(e, str)],
        }
    except Exception as exc:  # noqa: BLE001 — metadata is best-effort; never abort ingest
        llm_state["consecutive_failures"] = llm_state.get("consecutive_failures", 0) + 1
        if llm_state["consecutive_failures"] >= METADATA_LLM_MAX_CONSECUTIVE_FAILURES:
            llm_state["disabled"] = True
            print(
                f"[ingest]   ! metadata LLM disabled after "
                f"{METADATA_LLM_MAX_CONSECUTIVE_FAILURES} consecutive failures ({exc}) — "
                "regex fallback for the remaining articles",
                file=sys.stderr,
            )
        else:
            print(f"[ingest]   ! metadata LLM failed ({exc}) — using regex fallback", file=sys.stderr)
        return _regex_fallback_metadata(article_text)


# ---------------------------------------------------------------------------
# Chunk assembly
# ---------------------------------------------------------------------------

def _load_and_chunk(
    paths: list[Path],
    *,
    law_name: str | None,
    valid_from: str | None,
    llm_client,
) -> list[dict]:
    chunks: list[dict] = []
    llm_state: dict = {"consecutive_failures": 0, "disabled": False}
    for path in paths:
        print(f"[ingest] Reading {path.name} ...")
        full_text = _extract_text(path)

        # Legacy font-hack files (Latin glyphs standing in for Cyrillic) would
        # embed as garbage — skip them rather than poison retrieval.
        letters = re.findall(r"[A-Za-zЀ-ӿ]", full_text)
        cyrillic = re.findall(r"[Ѐ-ӿ]", full_text)
        if letters and (len(cyrillic) / len(letters)) < MIN_CYRILLIC_RATIO:
            print(
                f"[ingest]   ! SKIPPED {path.name}: only "
                f"{len(cyrillic) / len(letters):.0%} Cyrillic — legacy font encoding",
                file=sys.stderr,
            )
            continue

        # CLI flags override; otherwise identify the law per file from its first page
        if law_name and valid_from:
            doc_info = {"law": law_name, "valid_from": valid_from}
        else:
            detected = _detect_document_info(llm_client, full_text[:4000], fallback_name=path.stem)
            doc_info = {
                "law": law_name or detected["law"],
                "valid_from": valid_from or detected["valid_from"],
            }
        print(f"[ingest]   -> law: {doc_info['law']} | valid_from: {doc_info['valid_from']}")

        articles = _split_articles(full_text)
        print(f"[ingest]   -> {len(articles)} article segments")

        for article in articles:
            metadata = _extract_article_metadata(llm_client, article["text"], llm_state)
            parts = _subsplit_long_article(article["text"])
            for part_index, part_text in enumerate(parts):
                chunks.append({
                    "text": part_text,
                    "law": doc_info["law"],
                    "article": article["article"],
                    "title": metadata["title"],
                    "waste_types": metadata["waste_types"],
                    "entity_types": metadata["entity_types"],
                    "valid_from": doc_info["valid_from"],
                    "source_file": path.name,
                    "part": part_index if len(parts) > 1 else None,
                    "chunk_index": len(chunks),
                })
    return chunks


# ---------------------------------------------------------------------------
# Embedding (local multilingual-e5-large)
# ---------------------------------------------------------------------------

def _embed_batches(texts: list[str]) -> list[list[float]]:
    from sentence_transformers import SentenceTransformer

    print(f"[ingest] Loading {EMBED_MODEL_NAME} (first run downloads ~2 GB) ...")
    model = SentenceTransformer(EMBED_MODEL_NAME)

    prefixed = [E5_PASSAGE_PREFIX + t for t in texts]
    total_batches = -(-len(prefixed) // EMBED_BATCH_SIZE)
    vectors: list[list[float]] = []
    for i in range(0, len(prefixed), EMBED_BATCH_SIZE):
        batch = prefixed[i : i + EMBED_BATCH_SIZE]
        print(f"[ingest] Embedding batch {i // EMBED_BATCH_SIZE + 1}/{total_batches} ({len(batch)} chunks) ...")
        embeddings = model.encode(batch, normalize_embeddings=True, show_progress_bar=False)
        vectors.extend(vec.tolist() for vec in embeddings)
    return vectors


def _point_id(chunk: dict) -> str:
    """Content-addressed point id — the same chunk always resolves to the same id.

    Random uuid4 ids were what forced the old delete-the-whole-collection-first
    design: without a stable id a re-ingest could only append, doubling the
    corpus. Keyed on position (file/article/part) *and* a hash of the text, so
    edited text lands on a new id while untouched text overwrites itself.
    """
    text_hash = hashlib.sha256(chunk["text"].encode("utf-8")).hexdigest()[:CHUNK_HASH_PREFIX_LEN]
    key = f"{chunk['source_file']}|{chunk['article']}|{chunk['part']}|{text_hash}"
    return str(uuid.uuid5(POINT_ID_NAMESPACE, key))


def _upsert(
    client: QdrantClient,
    collection_name: str,
    chunks: list[dict],
    vectors: list[list[float]],
    run_id: str,
) -> None:
    points = [
        PointStruct(
            id=_point_id(chunks[i]),
            vector=vectors[i],
            payload={**chunks[i], "ingest_run_id": run_id},
        )
        for i in range(len(chunks))
    ]
    for i in range(0, len(points), UPSERT_BATCH_SIZE):
        client.upsert(collection_name=collection_name, points=points[i : i + UPSERT_BATCH_SIZE])
    print(f"[ingest] Upserted {len(points)} points into '{collection_name}'.")


def _prune_superseded(
    client: QdrantClient, collection_name: str, source_files: set[str], run_id: str
) -> None:
    """Drop points of the ingested files that this run did not rewrite.

    Deterministic ids mean the upsert already replaced every chunk that still
    exists, so whatever is left carrying an older run_id is a chunk the source
    document no longer contains (an article was deleted, or re-parsing split it
    differently). Scoped to the files just ingested — points belonging to any
    other document are never matched.

    Runs strictly after the upsert so there is no window in which a file has
    been cleared but not yet rewritten.
    """
    for source_file in sorted(source_files):
        client.delete(
            collection_name=collection_name,
            points_selector=FilterSelector(
                filter=Filter(
                    must=[FieldCondition(key="source_file", match=MatchValue(value=source_file))],
                    must_not=[
                        FieldCondition(key="ingest_run_id", match=MatchValue(value=run_id))
                    ],
                )
            ),
        )
    print(f"[ingest] Pruned superseded points for {len(source_files)} file(s).")


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Ingest law PDFs/DOCX into Qdrant.")
    parser.add_argument("files", nargs="+", help="PDF/DOCX files to ingest")
    parser.add_argument(
        "--recreate",
        action="store_true",
        help="drop and rebuild the whole collection (snapshots first); default is a non-destructive upsert",
    )
    # Optional: when omitted, each file's law title/date is auto-detected from
    # its first page (the corpus bundles many amendments with opaque filenames).
    parser.add_argument("--law-name", default=None, help='override, e.g. "Law on Waste Management 216/2021"')
    parser.add_argument("--valid-from", default=None, help="override ISO date the law/amendment took effect, e.g. 2021-01-01")
    parser.add_argument("--no-metadata-llm", action="store_true", help="skip LLM metadata extraction (regex fallback only)")
    args = parser.parse_args()

    if args.valid_from is not None:
        try:
            date.fromisoformat(args.valid_from)
        except ValueError:
            parser.error(f"--valid-from must be an ISO date (YYYY-MM-DD), got: {args.valid_from}")
    return args


def main() -> None:
    args = _parse_args()

    paths = [Path(p).resolve() for p in args.files]
    problems = [p for p in paths if not p.exists() or p.suffix.lower() not in SUPPORTED_EXTENSIONS]
    if problems:
        for p in problems:
            print(f"[ingest] Missing or unsupported file: {p}", file=sys.stderr)
        sys.exit(1)

    collection_name = os.getenv("QDRANT_COLLECTION", DEFAULT_COLLECTION)

    qdrant = _build_qdrant_client()
    llm_client = None if args.no_metadata_llm else _build_metadata_llm_client()
    if llm_client is None and not args.no_metadata_llm:
        print("[ingest] No GITHUB_MODELS_TOKEN — falling back to regex metadata.", file=sys.stderr)

    # Order matters more than anything else in this function. Extraction, OCR,
    # metadata and embedding all complete *before* the collection is touched, so
    # the failure modes that dominate this pipeline — OCR crash, LLM quota
    # exhausted, cancelled run — leave the existing corpus fully intact. The
    # previous order deleted first and wrote back last, turning any interruption
    # into total loss; that is how the corpus was destroyed twice.
    chunks = _load_and_chunk(paths, law_name=args.law_name, valid_from=args.valid_from, llm_client=llm_client)
    if not chunks:
        print("[ingest] No chunks extracted — aborting.", file=sys.stderr)
        sys.exit(1)

    print(f"[ingest] Total chunks: {len(chunks)}")
    vectors = _embed_batches([c["text"] for c in chunks])

    run_id = uuid.uuid4().hex
    _ensure_collection(qdrant, collection_name, recreate=args.recreate)
    _upsert(qdrant, collection_name, chunks, vectors, run_id)
    _prune_superseded(qdrant, collection_name, {c["source_file"] for c in chunks}, run_id)
    print("[ingest] Done.")


if __name__ == "__main__":
    main()
